from hashlib import sha256
from io import BytesIO
from pathlib import Path
import zipfile

from docx import Document

from services.docx_exporter import (
    ASD_TEMPLATE_PATH,
    PSYCHED_BOY_TEMPLATE_PATH,
    PSYCHED_GIRL_TEMPLATE_PATH,
    create_report_docx,
    has_branded_template,
)


def _hash(path: Path) -> str:
    return sha256(path.read_bytes()).hexdigest()


def test_asd_export_fills_metadata_and_preserves_template_structure():
    before_hash = _hash(ASD_TEMPLATE_PATH)
    exported = create_report_docx(
        title="ASD Clinical Diagnostic Assessment Report",
        patient_name="Patient Name",
        report_type="evaluation",
        content=(
            "# Report Metadata (Front Page)\n\n"
            "Report title: Test\n"
            "Client full name: Alex Example\n"
            "Date of birth: 2012-05-03"
        ),
    )

    output = Document(exported)
    assert _hash(ASD_TEMPLATE_PATH) == before_hash
    assert output.paragraphs[1].text == "Test"
    assert output.tables[0].cell(0, 0).text == "Name: Alex Example"
    assert output.tables[0].cell(0, 1).text == "Date of Birth: 2012-05-03"
    assert len(output.sections) == 1
    assert len(output.tables) == 11
    assert output.sections[0].different_first_page_header_footer is True


def test_asd_export_replaces_generated_section_body_only():
    exported = create_report_docx(
        title="ASD Clinical Diagnostic Assessment Report",
        patient_name="Alex Example",
        report_type="evaluation",
        content=(
            "# Reason for Referral\n\n"
            "Alex was referred for an autism assessment following concerns about social communication.\n\n"
            "# General Behavioural Observations\n\n"
            "Alex participated cooperatively and benefited from clear structure."
        ),
    )

    output = Document(exported)
    texts = [paragraph.text for paragraph in output.paragraphs]
    referral_index = texts.index("Reason for Referral")
    summary_index = texts.index("Summary & Conclusions")
    assert texts[referral_index + 1 : summary_index] == [
        "Alex was referred for an autism assessment following concerns about social communication."
    ]
    observations_index = texts.index("General Behavioural Observations")
    autism_index = texts.index("Autism Characteristics")
    assert texts[observations_index + 1 : autism_index] == [
        "Alex participated cooperatively and benefited from clear structure."
    ]
    assert "Highlighted Recommendations" in texts


def test_asd_export_retains_page_field_and_package_opens_as_docx():
    exported = create_report_docx(
        title="ASD Clinical Diagnostic Assessment Report",
        patient_name="Patient Name",
        report_type="evaluation",
        content="# Report Metadata (Front Page)\n\nReport title: Test",
    )

    with zipfile.ZipFile(BytesIO(exported.getvalue())) as archive:
        assert "word/document.xml" in archive.namelist()
        assert b"PAGE" in archive.read("word/footer1.xml")


def test_has_branded_template():
    assert has_branded_template("ASD Clinical Diagnostic Assessment Report") is True
    assert has_branded_template("SunnyHill CDBC Psychology Assessment - Boy") is True
    assert has_branded_template("Psycho-Educational Assessment - Girl") is True
    assert has_branded_template("Standard Intake Assessment") is False
    assert has_branded_template("Default Neuropsych Template") is False


def test_psyched_export_picks_boy_or_girl_template_by_title():
    before_boy = _hash(PSYCHED_BOY_TEMPLATE_PATH)
    before_girl = _hash(PSYCHED_GIRL_TEMPLATE_PATH)

    boy_export = create_report_docx(
        title="Psycho-Educational Assessment - Boy",
        patient_name="Ethan Chen",
        report_type="psychoeducational",
        content="# Report Metadata (Front Page)\n\nClient full name: Ethan Chen",
    )
    girl_export = create_report_docx(
        title="Psycho-Educational Assessment - Girl",
        patient_name="Ava Thompson",
        report_type="psychoeducational",
        content="# Report Metadata (Front Page)\n\nClient full name: Ava Thompson",
    )

    # Source templates are read-only inputs, never mutated.
    assert _hash(PSYCHED_BOY_TEMPLATE_PATH) == before_boy
    assert _hash(PSYCHED_GIRL_TEMPLATE_PATH) == before_girl

    boy_doc = Document(boy_export)
    girl_doc = Document(girl_export)
    boy_text = "\n".join(c.text for t in boy_doc.tables for r in t.rows for c in r.cells)
    girl_text = "\n".join(c.text for t in girl_doc.tables for r in t.rows for c in r.cells)
    assert "Ethan Chen" in boy_text
    assert "Ava Thompson" in girl_text


def test_psyched_export_fills_front_page_table_and_sections():
    exported = create_report_docx(
        title="Psycho-Educational Assessment - Boy",
        patient_name="Ethan Chen",
        report_type="psychoeducational",
        content=(
            "# Report Metadata (Front Page)\n\n"
            "Client full name: Ethan Chen\n"
            "Date of birth: 2016-04-02\n"
            "Examiner: Dr. A. Baillie\n\n"
            "# Reason for Referral\n\n"
            "Ethan was referred for a psycho-educational assessment."
        ),
    )

    output = Document(exported)
    front_cells = [c.text for r in output.tables[0].rows for c in r.cells if c.text.strip()]
    assert "Name: Ethan Chen" in front_cells
    assert "Date of Birth: 2016-04-02" in front_cells
    assert any(c.startswith("Examiner: Dr. A. Baillie") for c in front_cells)

    body_text = "\n".join(p.text for p in output.paragraphs)
    assert "Ethan was referred for a psycho-educational assessment." in body_text
    # The template's own placeholder name must not survive the export.
    assert "Joe" not in body_text
    assert "Lastname" not in body_text


def test_psyched_export_preserves_boilerplate_for_unfilled_sections():
    exported = create_report_docx(
        title="Psycho-Educational Assessment - Boy",
        patient_name="Ethan Chen",
        report_type="psychoeducational",
        content="# Report Metadata (Front Page)\n\nClient full name: Ethan Chen",
    )

    output = Document(exported)
    texts = [p.text for p in output.paragraphs]
    family_history_index = texts.index("Family History")
    # The section was never filled in, so the template's own instructional
    # boilerplate must still be there, untouched.
    assert "Where the child lives, who they live with." in texts[family_history_index + 1]


def test_generic_export_used_for_report_types_without_a_branded_template():
    exported = create_report_docx(
        title="Standard Intake Assessment",
        patient_name="Marcus Lee",
        report_type="evaluation",
        content=(
            "# Report Metadata (Front Page)\n\n"
            "Client full name: Marcus Lee\n"
            "Date of birth: 2014-02-11\n\n"
            "# Reason for Referral\n\n"
            "Marcus was referred for an intake assessment."
        ),
    )

    output = Document(exported)
    texts = [p.text for p in output.paragraphs]
    full_text = "\n".join(texts)

    assert texts[0] == "Standard Intake Assessment"
    assert "Marcus Lee" in full_text
    assert "2014-02-11" in full_text
    assert "Evaluation" in full_text  # report_type, humanized
    assert "confidential report" in full_text.lower()
    assert "Reason for Referral" in texts
    assert "Marcus was referred for an intake assessment." in full_text
    # The metadata block is surfaced as fields, not re-rendered as a section.
    assert "Report Metadata (Front Page)" not in texts


def test_generic_export_skips_empty_sections_and_opens_as_valid_docx():
    exported = create_report_docx(
        title="Default Neuropsych Template",
        patient_name="Patient Name",
        report_type="intake",
        content="# Reason for Referral\n\nReferred for evaluation.",
    )

    with zipfile.ZipFile(BytesIO(exported.getvalue())) as archive:
        assert "word/document.xml" in archive.namelist()

    output = Document(exported)
    assert any(p.text == "Reason for Referral" for p in output.paragraphs)
