"""
Document processing service for extracting text from various file formats
"""
from pathlib import Path
from typing import Optional
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
import textract
import pdfplumber


WORD_MIME_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
}


def is_word_doc(file_path: Path, content_type: Optional[str]) -> bool:
    if content_type in WORD_MIME_TYPES:
        return True

    return file_path.suffix.lower() in {".doc", ".docx"}

async def process_document(file_path: Path, content_type: Optional[str]) -> str:

    """
    Process a document and extract text content
    
    Args:
        file_path: Path to the document
        content_type: MIME type of the document
    
    Returns:
        Extracted text content
    """
    try:
        # For now, handle text files directly
        if content_type and "text" in content_type:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        
        # Handle PDF files
        elif content_type == "application/pdf":
            return await extract_pdf_text(file_path)
        
        # Handle Word documents (by MIME or extension)
        elif is_word_doc(file_path, content_type):
            return await extract_docx_text(file_path)
        
        # Default: try to read as text
        else:
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    return f.read()
            except:
                return f"[Binary file: {file_path.name}]"
    
    except Exception as e:
        return f"[Error processing file: {str(e)}]"

async def extract_pdf_text(file_path: Path) -> str:
    """Extract text from PDF using PyPDF2.

    Falls back to a helpful message if PyPDF2 is not installed.
    """
    if PdfReader is None:
        return f"[PDF extraction requires PyPDF2; please install backend/requirements.txt dependencies]"

    try:
        reader = PdfReader(str(file_path))
        texts: list[str] = []
        for page in reader.pages:
            # PyPDF2's page.extract_text() may be None for pages without extractable text
            text = page.extract_text()
            if text:
                texts.append(text)
        joined = "\n".join(texts).strip()
        return joined if joined else f"[No text found in PDF {file_path.name}]"
    except Exception as e:
        return f"[Error extracting PDF: {e}]"

async def extract_docx_text(file_path: Path) -> str:
    """Extract text from Word documents (.docx) using python-docx.

    - For `.docx` files this will return the concatenated paragraph text.
    - For legacy `.doc` files, returns a short message (conversion required).
    - If `python-docx` is not installed, returns a helpful message.
    """

    suffix = file_path.suffix.lower()
    if suffix == ".doc":
        if textract is None:
            return f"[Legacy .doc extraction requires textract and external dependencies (antiword/catdoc); please install backend/requirements.txt dependencies and system prerequisites]"
        try:
            raw = textract.process(str(file_path))
            if isinstance(raw, bytes):
                text = raw.decode("utf-8", errors="replace").strip()
            else:
                text = str(raw).strip()
            return text if text else f"[No text found in legacy .doc {file_path.name}]"
        except Exception as e:
            return f"[Error extracting legacy .doc: {e}]"

    if DocxDocument is None:
        return f"[DOCX extraction requires python-docx; please install backend/requirements.txt dependencies]"

    try:
        doc = DocxDocument(str(file_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text]
        joined = "\n".join(paragraphs).strip()
        return joined if joined else f"[No text found in Word document {file_path.name}]"
    except Exception as e:
        return f"[Error extracting Word document: {e}]"


async def extract_pdf_tables(file_path: Path) -> list:
    """Extract tables from a PDF using pdfplumber.

    Returns a list of tables; each table is a list of rows, each row is a list of cell strings.
    """
    if pdfplumber is None:
        return []

    try:
        tables_all: list[list[list[str]]] = []
        with pdfplumber.open(str(file_path)) as pdf:
            for page in pdf.pages:
                page_tables = page.extract_tables()
                for table in page_tables:
                    # table is a list of rows; each row is a list of cell values
                    # Normalize None to empty strings
                    normalized = [[cell if cell is not None else "" for cell in row] for row in table]
                    tables_all.append(normalized)
        return tables_all
    except Exception:
        return []


async def extract_docx_tables(file_path: Path) -> list:
    """Extract tables from a .docx using python-docx.

    Returns a list of tables; each table is a list of rows, each row is a list of cell strings.
    """
    if DocxDocument is None:
        return []

    try:
        doc = DocxDocument(str(file_path))
        tables_all: list[list[list[str]]] = []
        for table in doc.tables:
            rows: list[list[str]] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            tables_all.append(rows)
        return tables_all
    except Exception:
        return []


async def extract_tables_from_document(file_path: Path, content_type: Optional[str]) -> list:
    """High-level table extraction that dispatches based on MIME or extension.

    - PDF: uses `pdfplumber`.
    - DOCX: uses `python-docx`.
    - DOC: attempts to use `textract` but may lose structure.
    Returns a list of tables (possibly empty).
    """
    suffix = file_path.suffix.lower()
    if content_type == "application/pdf" or suffix == ".pdf":
        return await extract_pdf_tables(file_path)

    if suffix == ".docx" or (content_type in WORD_MIME_TYPES):
        return await extract_docx_tables(file_path)

    if suffix == ".doc":
        # textract will return plain text; reconstructing tables is non-trivial.
        if textract is None:
            return []
        try:
            raw = textract.process(str(file_path))
            if isinstance(raw, bytes):
                text = raw.decode("utf-8", errors="replace")
            else:
                text = str(raw)
            # As a fallback, return a single table with each line as a single-cell row
            rows = [[line] for line in text.splitlines() if line.strip()]
            return [rows] if rows else []
        except Exception:
            return []

    return []
