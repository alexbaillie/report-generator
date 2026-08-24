"""Template-preserving Word export for completed reports."""

from __future__ import annotations

from copy import deepcopy
from html import unescape
from io import BytesIO
from pathlib import Path
import re
from typing import Dict, Iterable, List, Mapping, Optional, Sequence, Tuple
from zipfile import ZIP_DEFLATED, ZipFile

from lxml import etree


_TEMPLATE_DIR = Path(__file__).resolve().parent.parent / "templates"
ASD_TEMPLATE_PATH = _TEMPLATE_DIR / "asd_school_age_boy_template.docx"
CDBC_BOY_TEMPLATE_PATH = _TEMPLATE_DIR / "sunnyhill_cdbc_boy_template.docx"
CDBC_GIRL_TEMPLATE_PATH = _TEMPLATE_DIR / "sunnyhill_cdbc_girl_template.docx"
PSYCHED_BOY_TEMPLATE_PATH = _TEMPLATE_DIR / "psyched_boy_template.docx"
PSYCHED_GIRL_TEMPLATE_PATH = _TEMPLATE_DIR / "psyched_girl_template.docx"

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
XML_NS = "http://www.w3.org/XML/1998/namespace"
NS = {"w": W_NS}


class DocxExportError(RuntimeError):
    """Raised when a report cannot be exported with a configured DOCX template."""


ASD_SECTION_TARGETS: Sequence[Tuple[str, Sequence[str]]] = (
    ("reason for referral", ("Reason for Referral",)),
    ("summary & conclusions", ("Summary & Conclusions",)),
    ("asd diagnostic criteria", ("ASD:",)),
    ("diagnostic considerations", ("All reports:",)),
    ("highlighted recommendations", ("Highlighted Recommendations",)),
    ("goals for the present assessment", ("Goals for the Present Assessment",)),
    ("presenting concerns", ("Presenting Concerns",)),
    ("family history", ("Family History",)),
    ("developmental & medical history", ("Developmental & Medical History",)),
    ("educational history", ("Educational History",)),
    ("previous assessments & interventions", ("Previous Assessments & Interventions",)),
    ("tests administered", ("Tests Administered",)),
    ("general behavioural observations", ("General Behavioural Observations",)),
    ("autism characteristics", ("Autism Characteristics",)),
    ("rating scales & adaptive behaviour", ("Adaptive Behaviour",)),
    ("additional recommendations", ("Additional Recommendations",)),
    ("appendices", ("APPENDIX 1: Interpretation of Test Results",)),
)


FRONT_TABLE_LABELS: Mapping[str, Sequence[str]] = {
    "name": ("client full name", "preferred name", "name"),
    "date of birth": ("date of birth",),
    "dates of evaluation": ("date(s) of evaluation", "dates of evaluation", "date(s) of assessment"),
    "date of evaluation": ("date(s) of evaluation", "dates of evaluation", "date(s) of assessment"),
    "date of assessment": ("date(s) of assessment", "dates of assessment", "date(s) of evaluation"),
    "date of report": ("date of report",),
    "chronological age": ("chronological age", "age at assessment"),
    "examiner": ("examiner", "assessing psychologist"),
    "copies to": ("copies to",),
}


# The Sunny Hill CDBC template uses ALL-CAPS/bold headings and a paragraph-based
# (not table) front page. Map each generated report section to the heading in the
# Word template whose following content it should replace.
CDBC_SECTION_TARGETS: Sequence[Tuple[str, Sequence[str]]] = (
    ("diagnostic summary (dsm-5-tr)", ("DIAGNOSTIC SUMMARY (DSM-5-TR)",)),
    ("summary of findings", ("SUMMARY OF FINDINGS",)),
    ("summary of assessment results", ("Summary of Assessment Results",)),
    ("highlighted recommendations", ("RECOMMENDATIONS",)),
    ("resources", ("RESOURCES",)),
    ("reason for referral", ("REASON FOR REFERRAL",)),
    ("sources of information", ("SOURCES OF INFORMATION",)),
    ("presenting concerns (as provided by parents)", ("Presenting Concerns as provided by parents",)),
    ("areas of relative strength & interests", ("Areas of Relative Strength and Interests",)),
    ("family history", ("Family History",)),
    ("developmental & medical history", ("Developmental & Medical History",)),
    ("educational history", ("Educational History",)),
    ("previous assessments & interventions", ("Previous Assessments & Interventions",)),
    ("behavioural observations", ("BEHAVIOURAL OBSERVATIONS",)),
    ("tests administered", ("TESTS ADMINISTERED",)),
    ("cognitive, memory & academic results", ("General Cognitive Abilities",)),
)


# The Psycho-Educational template has a table-based front page (same shape as
# ASD) and mostly-exact heading matches; a few sections cover several
# subsections in the Word doc (e.g. all the cognitive score tables between
# "Test Results" and "Academic Abilities") — anchoring on the first heading in
# that span means the generated content replaces the whole span, consistent
# with how ASD/CDBC handle "Recommendations"-style catch-all sections.
PSYCHED_SECTION_TARGETS: Sequence[Tuple[str, Sequence[str]]] = (
    ("reason for referral", ("Reason for Referral",)),
    ("presenting concerns (as reported by parents)", ("Presenting Concerns as Reported by Parents",)),
    ("teacher's perspective", ("Teacher's Perspective",)),
    ("child's perspective", ("Joe's Perspective", "Jane's Perspective")),
    ("areas of relative strength & interests", ("Areas of Relative Strength & Interests",)),
    ("family history", ("Family History",)),
    ("developmental & medical history", ("Developmental & Medical History",)),
    ("educational history", ("Educational History",)),
    ("previous assessments & interventions", ("Previous Assessments & Interventions",)),
    ("tests administered", ("Tests Administered",)),
    ("behavioural observations", ("Behavioural Observations",)),
    ("cognitive & neuropsychological results", ("Test Results",)),
    ("academic achievement", ("Academic Abilities",)),
    ("adaptive, behavioural & social-emotional functioning", ("Adaptive Behaviour",)),
    ("summary of assessment results", ("Summary of Assessment Results",)),
    ("diagnoses (dsm-5-tr)", ("Descriptors/Diagnoses (DSM-5-TR)",)),
    ("highlighted recommendations", ("Highlighted Recommendations",)),
    ("recommendations", ("Recommendations",)),
)


# Front-page paragraphs like "DATE OF BIRTH:" get the metadata value appended.
CDBC_FRONT_LABELS: Mapping[str, Sequence[str]] = {
    "chart number": ("chart number",),
    "date of birth": ("date of birth",),
    "date of assessment": ("date(s) of assessment", "dates of assessment", "date of assessment"),
    "age at assessment": ("age at assessment", "chronological age"),
    "date of report": ("date of report",),
    "date of conference": ("date of conference",),
}


def supports_asd_template(title: str) -> bool:
    normalized = _normalize(title)
    return "asd" in normalized or "autism" in normalized


def supports_cdbc_template(title: str) -> bool:
    normalized = _normalize(title)
    return "cdbc" in normalized or "sunny hill" in normalized or "sunnyhill" in normalized


def supports_psyched_template(title: str) -> bool:
    normalized = _normalize(title)
    return "psycho-educational" in normalized or "psychoeducational" in normalized or "psyched" in normalized


def has_branded_template(title: str) -> bool:
    """True if this title matches a real branded Word template (ASD/CDBC/PsychEd).
    Every report is still exportable via the generic fallback below — this only
    tells you whether the output matches a clinic's specific letterhead/layout."""
    return supports_asd_template(title) or supports_cdbc_template(title) or supports_psyched_template(title)


def _select_profile(title: str):
    """Return (template_path, section_targets, front_page_filler) for the title,
    or None if there's no branded template — the caller falls back to a plain,
    generic export in that case."""
    if supports_cdbc_template(title):
        template = CDBC_GIRL_TEMPLATE_PATH if "girl" in _normalize(title) else CDBC_BOY_TEMPLATE_PATH
        return template, CDBC_SECTION_TARGETS, _fill_front_page_cdbc
    if supports_psyched_template(title):
        template = PSYCHED_GIRL_TEMPLATE_PATH if "girl" in _normalize(title) else PSYCHED_BOY_TEMPLATE_PATH
        return template, PSYCHED_SECTION_TARGETS, _fill_front_page_psyched
    if supports_asd_template(title):
        return ASD_TEMPLATE_PATH, ASD_SECTION_TARGETS, _fill_front_page_asd
    return None


def export_filename(title: str) -> str:
    safe = re.sub(r"[^A-Za-z0-9._-]+", "_", title).strip("._")
    return f"{safe or 'report'}.docx"


def parse_markdown_sections(content: str) -> Dict[str, str]:
    sections: Dict[str, List[str]] = {}
    current: Optional[str] = None

    for raw_line in (content or "").replace("\r\n", "\n").split("\n"):
        match = re.match(r"^\s*#{1,6}\s+(.+?)\s*$", raw_line)
        if match:
            current = _normalize(match.group(1))
            sections.setdefault(current, [])
            continue
        if current is not None:
            sections[current].append(raw_line)

    if not sections and content.strip():
        sections["report"] = [content.strip()]

    return {
        name: "\n".join(lines).strip()
        for name, lines in sections.items()
        if "\n".join(lines).strip()
    }


def parse_markdown_sections_ordered(content: str) -> List[Tuple[str, str]]:
    """Like parse_markdown_sections, but keeps the original heading text and
    document order — used for the generic export, which renders real headings
    rather than looking sections up by normalized key."""
    sections: List[Tuple[str, List[str]]] = []
    for raw_line in (content or "").replace("\r\n", "\n").split("\n"):
        match = re.match(r"^\s*#{1,6}\s+(.+?)\s*$", raw_line)
        if match:
            sections.append((match.group(1).strip(), []))
            continue
        if sections:
            sections[-1][1].append(raw_line)

    if not sections and content.strip():
        return [("Report", [content.strip()])]

    return [
        (heading, "\n".join(lines).strip())
        for heading, lines in sections
        if "\n".join(lines).strip()
    ]


def create_report_docx(
    *,
    title: str,
    patient_name: str,
    report_type: str,
    content: str,
) -> BytesIO:
    profile = _select_profile(title)
    if profile is None:
        return _create_generic_report_docx(
            title=title, patient_name=patient_name, report_type=report_type, content=content
        )
    template_path, section_targets, fill_front_page = profile
    if not template_path.is_file():
        raise DocxExportError("The Word template for this report type is missing from the application.")

    with ZipFile(template_path, "r") as source:
        document_xml = source.read("word/document.xml")

    root = etree.fromstring(document_xml)
    sections = parse_markdown_sections(content)
    metadata = _extract_labeled_values(sections.get("report metadata (front page)", ""))

    fill_front_page(root, title=title, patient_name=patient_name, metadata=metadata)
    _fill_generated_sections(root, sections, section_targets)

    # Make the leftover template boilerplate readable: use the real name instead
    # of the "Jane/Joe Lastname" placeholders, and drop instructional highlighting.
    clean_name = _clean_name(metadata.get("client full name") or metadata.get("name") or patient_name)
    _replace_placeholder_names(root, clean_name)
    _strip_highlighting(root)

    patched_xml = etree.tostring(
        root,
        encoding="UTF-8",
        xml_declaration=True,
        standalone=True,
    )
    return _write_patched_package(patched_xml, template_path)


def _create_generic_report_docx(
    *,
    title: str,
    patient_name: str,
    report_type: str,
    content: str,
) -> BytesIO:
    """Plain, cleanly-formatted export for report types with no branded Word
    template (e.g. Standard Intake, Neuropsych) — a real heading per section
    instead of the clinic's specific letterhead/layout. Every report type is
    exportable this way; only some also get the branded version."""
    from docx import Document as _Document  # local import: only needed here

    document = _Document()
    ordered_sections = parse_markdown_sections_ordered(content)

    document.add_heading(title or "Psychological Report", level=0)

    info = document.add_paragraph()
    info.add_run(f"Patient: {patient_name or 'Unknown'}").bold = True
    if report_type:
        document.add_paragraph(f"Report type: {report_type.replace('_', ' ').title()}")

    # The front-page metadata section is structured "Label: value" lines (the
    # same convention the branded exporters read) — surface it as a compact
    # field list rather than rendering it as a regular narrative section.
    metadata_heading = next(
        (h for h, _ in ordered_sections if _normalize(h) in {"report metadata (front page)", "front page"}),
        None,
    )
    if metadata_heading:
        metadata_content = next(c for h, c in ordered_sections if h == metadata_heading)
        for label, value in _extract_labeled_values(metadata_content).items():
            if label in {"report title", "confidentiality statement"}:
                continue
            document.add_paragraph(f"{label.title()}: {value}")

    document.add_paragraph(
        "This is a confidential report. AI-generated content is a draft and must be "
        "reviewed by a licensed clinician before use."
    )

    for heading, section_content in ordered_sections:
        if heading == metadata_heading:
            continue
        document.add_heading(heading, level=1)
        for block in _content_blocks(section_content):
            document.add_paragraph(block)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def _write_patched_package(document_xml: bytes, template_path: Path) -> BytesIO:
    output = BytesIO()
    with ZipFile(template_path, "r") as source, ZipFile(
        output, "w", compression=ZIP_DEFLATED
    ) as destination:
        for info in source.infolist():
            payload = document_xml if info.filename == "word/document.xml" else source.read(info.filename)
            destination.writestr(info, payload)
    output.seek(0)
    return output


def _fill_front_page_table_based(
    root,
    *,
    title: str,
    patient_name: str,
    metadata: Mapping[str, str],
    title_anchor: str,
) -> None:
    """Shared front-page filler for templates with a title paragraph, a
    confidentiality sentence, and a labelled front-page table (ASD, PsychEd)."""
    paragraphs = root.xpath("/w:document/w:body/w:p", namespaces=NS)
    report_title = metadata.get("report title") or title
    title_paragraph = _find_paragraph(paragraphs, (title_anchor,))
    if title_paragraph is not None:
        _set_element_text(title_paragraph, report_title)

    confidentiality = metadata.get("confidentiality statement")
    if confidentiality:
        for paragraph in paragraphs:
            if _normalize(_element_text(paragraph)).startswith("this is a confidential report"):
                _set_element_text(paragraph, confidentiality)
                break

    tables = root.xpath("/w:document/w:body/w:tbl", namespaces=NS)
    if not tables:
        return

    effective_metadata = dict(metadata)
    if patient_name and _normalize(patient_name) not in {"", "patient name"}:
        effective_metadata.setdefault("client full name", patient_name)

    for cell in tables[0].xpath(".//w:tc", namespaces=NS):
        original = _element_text(cell).strip()
        if ":" not in original:
            continue
        source_label = _normalize(original.split(":", 1)[0])
        aliases = FRONT_TABLE_LABELS.get(source_label)
        if not aliases:
            continue
        value = _first_value(effective_metadata, aliases)
        if not value:
            continue
        label_text = original.split(":", 1)[0].strip()
        _set_cell_text(cell, f"{label_text}: {value}")


def _fill_front_page_asd(
    root,
    *,
    title: str,
    patient_name: str,
    metadata: Mapping[str, str],
) -> None:
    _fill_front_page_table_based(
        root,
        title=title,
        patient_name=patient_name,
        metadata=metadata,
        title_anchor="Clinical Diagnostic Assessment Report",
    )


def _fill_front_page_psyched(
    root,
    *,
    title: str,
    patient_name: str,
    metadata: Mapping[str, str],
) -> None:
    _fill_front_page_table_based(
        root,
        title=title,
        patient_name=patient_name,
        metadata=metadata,
        title_anchor="Psycho-Educational Assessment Report",
    )


def _fill_front_page_cdbc(
    root,
    *,
    title: str,
    patient_name: str,
    metadata: Mapping[str, str],
) -> None:
    paragraphs = root.xpath("/w:document/w:body/w:p", namespaces=NS)

    report_title = metadata.get("report title")
    if report_title:
        for paragraph in paragraphs:
            if _normalize(_element_text(paragraph)) == "psychology assessment report":
                _set_element_text(paragraph, report_title)
                break

    # The client's name is a stand-alone ALL-CAPS paragraph like "JANE LASTNAME".
    name = metadata.get("client full name") or metadata.get("name")
    if (not name) and patient_name and _normalize(patient_name) not in {"", "patient name"}:
        name = patient_name
    if name:
        for paragraph in paragraphs:
            normalized = _normalize(_element_text(paragraph))
            if normalized.endswith("lastname") and ":" not in normalized:
                _set_element_text(paragraph, name.upper())
                break
        # Swap the placeholder name inside the confidentiality sentence too.
        for paragraph in paragraphs:
            text = _element_text(paragraph)
            if _normalize(text).startswith("this is a confidential report"):
                swapped = re.sub(r"(?i)\b(?:jane|joe)\s+lastname\b", name, text)
                if swapped != text:
                    _set_element_text(paragraph, swapped)
                break

    # Labelled front-page paragraphs ("DATE OF BIRTH:") get their value appended.
    for paragraph in paragraphs:
        text = _element_text(paragraph).strip()
        if ":" not in text:
            continue
        source_label = _normalize(text.split(":", 1)[0])
        aliases = CDBC_FRONT_LABELS.get(source_label)
        if not aliases:
            continue
        value = _first_value(metadata, aliases)
        if not value:
            continue
        label_text = text.split(":", 1)[0].strip()
        _set_element_text(paragraph, f"{label_text}: {value}")


def _fill_generated_sections(root, sections: Mapping[str, str], section_targets) -> None:
    body = root.find(f"{{{W_NS}}}body")
    if body is None:
        raise DocxExportError("The Word template has no document body.")

    paragraphs = body.xpath("./w:p", namespaces=NS)
    located: List[Tuple[int, str, object]] = []
    children = list(body)

    for section_name, aliases in section_targets:
        paragraph = _find_paragraph(paragraphs, aliases)
        if paragraph is not None:
            located.append((children.index(paragraph), section_name, paragraph))

    located.sort(key=lambda item: item[0])
    for position, (_, section_name, heading) in enumerate(located):
        generated = sections.get(section_name)
        if not generated:
            continue

        current_children = list(body)
        start_index = current_children.index(heading)
        if position + 1 < len(located):
            end_index = current_children.index(located[position + 1][2])
        else:
            end_index = len(current_children)

        interval = current_children[start_index + 1 : end_index]
        style_source = next(
            (
                element
                for element in interval
                if element.tag == f"{{{W_NS}}}p" and _element_text(element).strip()
            ),
            heading,
        )
        for element in interval:
            if element.tag == f"{{{W_NS}}}p":
                body.remove(element)

        anchor = heading
        for block in _content_blocks(generated):
            new_paragraph = _paragraph_element(style_source, block)
            anchor.addnext(new_paragraph)
            anchor = new_paragraph


def _clean_name(name: str) -> str:
    """Patient name for substitution — parentheticals like "(test)" removed."""
    name = re.sub(r"\s*\(.*?\)\s*", " ", name or "").strip()
    if _normalize(name) in {"", "patient name"}:
        return ""
    return name


def _replace_placeholder_names(root, full_name: str) -> None:
    """Swap the template's "Jane/Joe Lastname" placeholders for the real name
    throughout the document, preserving ALL-CAPS styling where used."""
    if not full_name:
        return
    parts = full_name.split()
    first, last = parts[0], (parts[-1] if len(parts) > 1 else "")

    def _cased(match, replacement):
        return replacement.upper() if match.group(0).isupper() else replacement

    for node in root.iter(f"{{{W_NS}}}t"):
        if not node.text:
            continue
        text = re.sub(r"(?i)\b(?:jane|joe)\s+lastname\b", lambda m: _cased(m, full_name), node.text)
        text = re.sub(r"(?i)\b(?:jane|joe)\b", lambda m: _cased(m, first), text)
        if last:
            text = re.sub(r"(?i)\blastname\b", lambda m: _cased(m, last), text)
        if text != node.text:
            node.text = text


def _strip_highlighting(root) -> None:
    """Remove text highlighting and run-level shading (the yellow placeholder marks)."""
    for highlight in list(root.iter(f"{{{W_NS}}}highlight")):
        parent = highlight.getparent()
        if parent is not None:
            parent.remove(highlight)
    for shading in list(root.iter(f"{{{W_NS}}}shd")):
        parent = shading.getparent()
        if parent is not None and parent.tag == f"{{{W_NS}}}rPr":
            parent.remove(shading)


def _extract_labeled_values(text: str) -> Dict[str, str]:
    values: Dict[str, str] = {}
    for line in (text or "").splitlines():
        match = re.match(r"^\s*([^:#]{2,80}?):\s*(.+?)\s*$", line)
        if match and match.group(2).strip():
            values[_normalize(match.group(1))] = match.group(2).strip()
    return values


def _first_value(values: Mapping[str, str], aliases: Iterable[str]) -> Optional[str]:
    for alias in aliases:
        value = values.get(_normalize(alias))
        if value:
            return value
    return None


def _find_paragraph(paragraphs: Iterable, aliases: Iterable[str]):
    expected = {_normalize(alias) for alias in aliases}
    for paragraph in paragraphs:
        if _normalize(_element_text(paragraph)) in expected:
            return paragraph
    return None


def _element_text(element) -> str:
    return "".join(element.xpath(".//w:t/text()", namespaces=NS))


def _set_element_text(element, text: str) -> None:
    paragraph_properties = element.find(f"{{{W_NS}}}pPr")
    run_properties = element.find(f".//{{{W_NS}}}rPr")
    for child in list(element):
        if child is not paragraph_properties:
            element.remove(child)

    run = etree.Element(f"{{{W_NS}}}r")
    if run_properties is not None:
        run.append(deepcopy(run_properties))
    text_element = etree.SubElement(run, f"{{{W_NS}}}t")
    if text.startswith(" ") or text.endswith(" "):
        text_element.set(f"{{{XML_NS}}}space", "preserve")
    text_element.text = text
    element.append(run)


def _set_cell_text(cell, text: str) -> None:
    paragraphs = cell.xpath("./w:p", namespaces=NS)
    if paragraphs:
        _set_element_text(paragraphs[0], text)
        for paragraph in paragraphs[1:]:
            cell.remove(paragraph)
        return

    paragraph = etree.SubElement(cell, f"{{{W_NS}}}p")
    _set_element_text(paragraph, text)


def _paragraph_element(style_source, text: str):
    paragraph = etree.Element(f"{{{W_NS}}}p")
    paragraph_properties = style_source.find(f"{{{W_NS}}}pPr")
    if paragraph_properties is not None:
        paragraph.append(deepcopy(paragraph_properties))

    run = etree.SubElement(paragraph, f"{{{W_NS}}}r")
    run_properties = style_source.find(f".//{{{W_NS}}}rPr")
    if run_properties is not None:
        run.append(deepcopy(run_properties))
    text_element = etree.SubElement(run, f"{{{W_NS}}}t")
    if text.startswith(" ") or text.endswith(" "):
        text_element.set(f"{{{XML_NS}}}space", "preserve")
    text_element.text = text
    return paragraph


def _content_blocks(text: str) -> List[str]:
    cleaned = unescape(text or "")
    cleaned = re.sub(r"(?i)<br\s*/?>", "\n", cleaned)
    cleaned = re.sub(r"(?i)</(?:p|div|tr|li)>", "\n", cleaned)
    cleaned = re.sub(r"<[^>]+>", "", cleaned)
    blocks = [
        re.sub(r"\s+", " ", block).strip()
        for block in re.split(r"\n\s*\n|\n", cleaned)
    ]
    return [block for block in blocks if block]


def _normalize(value: str) -> str:
    folded = (value or "").replace("’", "'").replace("‘", "'")
    return re.sub(r"\s+", " ", folded.strip().casefold())
